"""
Chargement et découpage des documents pour le RAG fiction.
Version 2.0 avec support PDF et DOCX.
"""
from pathlib import Path
from typing import List, Optional
from langchain_community.document_loaders import TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document


# Extensions supportées par type
SUPPORTED_EXTENSIONS = {
    "text": [".txt", ".md"],
    "pdf": [".pdf"],
    "docx": [".docx", ".doc"],
}

ALL_EXTENSIONS = [ext for exts in SUPPORTED_EXTENSIONS.values() for ext in exts]


def load_text_file(file_path: Path) -> List[Document]:
    """
    Charge un fichier texte (.txt, .md).
    
    Args:
        file_path: Chemin vers le fichier
        
    Returns:
        Liste de documents LangChain
    """
    try:
        loader = TextLoader(str(file_path), encoding="utf-8")
        return loader.load()
    except UnicodeDecodeError:
        # Essayer avec d'autres encodages
        for encoding in ["latin-1", "cp1252", "iso-8859-1"]:
            try:
                loader = TextLoader(str(file_path), encoding=encoding)
                docs = loader.load()
                print(f"  ℹ️  Fichier chargé avec encodage {encoding}")
                return docs
            except UnicodeDecodeError:
                continue
        raise


def load_pdf_file(file_path: Path) -> List[Document]:
    """
    Charge un fichier PDF.
    
    Args:
        file_path: Chemin vers le fichier PDF
        
    Returns:
        Liste de documents LangChain
    """
    try:
        from langchain_community.document_loaders import PyPDFLoader
        loader = PyPDFLoader(str(file_path))
        return loader.load()
    except ImportError:
        # Fallback sur pypdf directement
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(file_path))
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return [Document(
                page_content=text,
                metadata={"source": str(file_path)}
            )]
        except ImportError:
            raise ImportError(
                "Pour charger des PDF, installez pypdf:\n"
                "pip install pypdf"
            )


def load_docx_file(file_path: Path) -> List[Document]:
    """
    Charge un fichier DOCX.
    
    Args:
        file_path: Chemin vers le fichier DOCX
        
    Returns:
        Liste de documents LangChain
    """
    try:
        from langchain_community.document_loaders import Docx2txtLoader
        loader = Docx2txtLoader(str(file_path))
        return loader.load()
    except ImportError:
        # Fallback sur python-docx
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(file_path))
            text = "\n".join([para.text for para in doc.paragraphs])
            return [Document(
                page_content=text,
                metadata={"source": str(file_path)}
            )]
        except ImportError:
            raise ImportError(
                "Pour charger des DOCX, installez python-docx:\n"
                "pip install python-docx"
            )


def load_document(file_path: Path) -> List[Document]:
    """
    Charge un document selon son extension.
    
    Args:
        file_path: Chemin vers le fichier
        
    Returns:
        Liste de documents LangChain
    """
    suffix = file_path.suffix.lower()
    
    if suffix in SUPPORTED_EXTENSIONS["text"]:
        return load_text_file(file_path)
    elif suffix in SUPPORTED_EXTENSIONS["pdf"]:
        return load_pdf_file(file_path)
    elif suffix in SUPPORTED_EXTENSIONS["docx"]:
        return load_docx_file(file_path)
    else:
        raise ValueError(f"Extension non supportée: {suffix}")


def load_project_documents(
    project_path: Path,
    extensions: Optional[List[str]] = None,
    verbose: bool = True
) -> List[Document]:
    """
    Charge tous les fichiers d'un projet.
    
    Args:
        project_path: Chemin vers le dossier du projet
        extensions: Liste d'extensions à charger (défaut: toutes supportées)
        verbose: Afficher la progression
        
    Returns:
        Liste de documents LangChain
    """
    if extensions is None:
        extensions = ALL_EXTENSIONS
    
    docs = []
    
    if not project_path.exists():
        raise FileNotFoundError(f"Le projet {project_path} n'existe pas")
    
    # Compter les fichiers d'abord
    files_to_load = []
    for path in project_path.rglob("*"):
        if path.is_file() and path.suffix.lower() in extensions:
            files_to_load.append(path)
    
    if verbose:
        print(f"📁 {len(files_to_load)} fichiers trouvés")
    
    # Charger chaque fichier
    for path in files_to_load:
        try:
            loaded_docs = load_document(path)
            
            # Enrichir les métadonnées
            for doc in loaded_docs:
                doc.metadata["relative_path"] = str(path.relative_to(project_path))
                doc.metadata["file_name"] = path.name
                doc.metadata["file_type"] = path.suffix.lower()
                doc.metadata["folder"] = path.parent.name
            
            docs.extend(loaded_docs)
            
            if verbose:
                print(f"✓ Chargé: {path.relative_to(project_path)}")
                
        except Exception as e:
            if verbose:
                print(f"✗ Erreur {path.name}: {e}")
    
    return docs


def split_documents(
    docs: List[Document],
    chunk_size: int = 1000,
    chunk_overlap: int = 150,
    separators: Optional[List[str]] = None
) -> List[Document]:
    """
    Découpe les documents en chunks pour l'embedding.
    
    Args:
        docs: Liste de documents à découper
        chunk_size: Taille maximale d'un chunk en caractères
        chunk_overlap: Chevauchement entre chunks
        separators: Séparateurs personnalisés
        
    Returns:
        Liste de chunks
    """
    if separators is None:
        # Séparateurs optimisés pour la fiction
        separators = [
            "\n\n\n",  # Changement de scène/chapitre
            "\n\n",    # Nouveau paragraphe
            "\n",      # Nouvelle ligne
            ".",       # Fin de phrase
            "!",       # Exclamation
            "?",       # Question
            ";",       # Point-virgule
            ",",       # Virgule
            " ",       # Espace
            ""         # Caractère
        ]
    
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=separators,
        length_function=len,
    )
    
    chunks = splitter.split_documents(docs)
    
    # Ajouter un index de chunk dans les métadonnées
    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk_index"] = i
    
    print(f"✓ {len(docs)} documents découpés en {len(chunks)} chunks")
    
    return chunks


def smart_split_documents(
    docs: List[Document],
    min_chunk_size: int = 500,
    max_chunk_size: int = 1500,
    overlap_ratio: float = 0.15
) -> List[Document]:
    """
    Découpage intelligent adaptatif selon le contenu.
    
    Adapte la taille des chunks selon le type de contenu:
    - Dialogues: chunks plus petits pour garder le contexte
    - Descriptions: chunks plus grands pour la cohérence
    
    Args:
        docs: Documents à découper
        min_chunk_size: Taille minimum
        max_chunk_size: Taille maximum
        overlap_ratio: Ratio de chevauchement
        
    Returns:
        Liste de chunks
    """
    all_chunks = []
    
    for doc in docs:
        content = doc.page_content
        
        # Estimer le type de contenu
        dialogue_ratio = content.count('"') / max(len(content), 1)
        
        # Adapter la taille
        if dialogue_ratio > 0.05:  # Beaucoup de dialogues
            chunk_size = min_chunk_size
        else:
            chunk_size = max_chunk_size
        
        chunk_overlap = int(chunk_size * overlap_ratio)
        
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )
        
        chunks = splitter.split_documents([doc])
        
        # Propager les métadonnées
        for chunk in chunks:
            chunk.metadata["adaptive_chunk_size"] = chunk_size
        
        all_chunks.extend(chunks)
    
    return all_chunks


# Fonctions utilitaires
def get_supported_extensions() -> List[str]:
    """Retourne la liste des extensions supportées."""
    return ALL_EXTENSIONS


def is_supported_file(file_path: Path) -> bool:
    """Vérifie si un fichier est supporté."""
    return file_path.suffix.lower() in ALL_EXTENSIONS


# Test du module
if __name__ == "__main__":
    import sys
    
    project = sys.argv[1] if len(sys.argv) > 1 else "anomalie2084"
    project_path = Path("data") / project
    
    print(f"\n📚 Test du chargement pour '{project}'")
    print(f"   Extensions supportées: {', '.join(ALL_EXTENSIONS)}")
    print("=" * 50)
    
    if project_path.exists():
        docs = load_project_documents(project_path)
        print(f"\n📊 Résultat: {len(docs)} documents chargés")
        
        if docs:
            chunks = split_documents(docs)
            print(f"   → {len(chunks)} chunks créés")
            
            # Afficher un aperçu
            print(f"\n📝 Aperçu du premier chunk:")
            print(f"   Source: {chunks[0].metadata.get('relative_path', 'inconnu')}")
            print(f"   Contenu: {chunks[0].page_content[:100]}...")
    else:
        print(f"❌ Projet non trouvé: {project_path}")
